"""Adds an introduction section to a Word document."""
import docx

from django.conf import settings


def introduction(content, document):
    """Add introduction section."""
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).width = docx.shared.Cm(12)

    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).add_paragraph(
        content.get('name'),
        style='Heading 1')
    table.cell(0, 0).add_paragraph(
        content.get('position'),
        style='Heading 2')

    table.cell(0, 1).width = docx.shared.Cm(4)
    table.cell(0, 1).add_picture((
        "/srv/rsum/static/{directory}/img/mockup/avatar-02.png").format(
            directory=settings.rsum.xander.DIR))
    table.cell(0, 1).paragraphs[0].paragraph_format.alignment = (
        docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    )
    return document
