"""Module for education."""
# pylint: disable=no-name-in-module,no-member
from django.conf import settings as django_settings
from docx.shared import Cm


class Education(object):
    """Education export object.

    .. attribute:: name

    .. attribute:: document

    .. attribute:: settings

    .. attribute:: line_spacing
    """

    name = None
    document = None
    settings = django_settings
    line_spacing = 1.0

    def save(self, section, document, graphics):
        """Save the education section of the document.

        :param: section
        :param: document
        :param: graphics
        :return: Updated document,
        """
        if graphics:
            document = self.get_education_graphics(section, document)
        else:
            document = self.get_education(section, document)
        return document

    def get_education(self, section, document):
        """Add education section.

        :param [dict(str, str)] education:
            Education section for current document.
        :param object document: Current document.
        :return: Current document with Educaiton section.
        :rtype: object
        """
        self.document = document
        education = section
        paragraph = document.add_paragraph(
            'Education',
            style='Heading 3')
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = 0
        paragraph = document.add_paragraph(
            education.get('name'),
            style='Heading 4')
        paragraph.paragraph_format.space_before = 0
        paragraph = document.add_paragraph(
            education.get('studies'),
            style='Heading 5')
        paragraph.paragraph_format.space_before = 0
        paragraph = document.add_paragraph(
            "{0}, {1}".format(
                education.get('location'),
                education.get('duration')),
            style='Heading 6')
        paragraph.paragraph_format.space_before = 0
        return document

    def get_education_graphics(self, section, document):
        """Add education section.
        :param [dict(str, str)] education:
            Education section for current document.
        :param object document: Current document.
        :return: Current document with Educaiton section.
        :rtype: object
        """
        self.document = document
        education = section
        paragraph = document.add_paragraph(
            'Education',
            style='Heading 3')
        paragraph.paragraph_format.line_spacing = self.line_spacing
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.page_break_before = True
        document = self.set_education_picture(document)
        paragraph = document.add_paragraph(
            education.get('name'),
            style='Heading 4')
        paragraph.paragraph_format.space_before = 0
        paragraph = document.add_paragraph(
            education.get('studies'),
            style='Heading 5')
        paragraph.paragraph_format.space_before = 0
        paragraph = document.add_paragraph(
            "{0}, {1}".format(
                education.get('location'),
                education.get('duration')),
            style='Heading 6')
        paragraph.paragraph_format.space_before = 0
        return document

    def set_education_picture(self, document):
        """Sets the picture for display in the education section.

        :param docx.Document document: The document to update.
        """
        document.add_picture(
            'static/profiles/{0}/img/1920x1080/01.jpg'.format(
                self.settings.DIR),
            width=Cm(4))
        return document
