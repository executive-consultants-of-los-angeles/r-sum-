"""Experience module."""
# pylint: disable=no-member
from django.conf import settings as django_settings
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


class Experience(object):
    """Export Experience object."""

    name = None
    cm = Cm
    projects = None
    settings = django_settings
    experience = None
    spacing = 0.9

    def save(self, name, section, document, graphics=True):
        """Save the experience section for dullards.

        :rtype: object
        """
        self.experience = section
        self.name = name

        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0
        paragraph.paragraph_format.page_break_before = True

        paragraph = document.add_paragraph(
            'Experience', style='Heading 3')

        document = self.add_intro(document)

        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, value in enumerate(self.experience):
            if index % 3 == 0:
                table.add_row()
            document = self.set_tables(
                document, table=table,
                value=value, index=index,
                graphics=False
            )

        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0
        return document

    def save_with_graphics(self, name, section, document):
        """Add experience section.

        :param [dict(str, str)] experience:
            Experience section to add to document.
        :param objectable document: Documentable to update.
        :return: Documentable updated with Experience section.
        :rtype: object
        """
        self.experience = section
        self.name = name

        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0
        paragraph.paragraph_format.page_break_before = True
        paragraph = document.add_paragraph(
            'Experience', style='Heading 3')

        document = self.add_intro(document)

        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, value in enumerate(self.experience):
            if index % 3 == 0:
                table.add_row()
            document = self.set_tables(
                document, table=table,
                value=value, index=index,
                graphics=True
            )

        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0
        return document

    def add_intro(self, document):
        """Add an introduction to the experience section."""
        introduction = self.experience.pop(0)
        paragraph = document.add_paragraph(
            introduction.get('introduction'), style='Normal')
        paragraph.paragraph_format.line_spacing = 0.7
        return document

    def set_tables(self, document, **dargs):
        """Set tables for the experience section."""
        document = (
            self.prep_tables(
                document, table=dargs.get('table'),
                value=dargs.get('value'), index=dargs.get('index'),
                graphics=dargs.get('graphics')
            )
        )
        return document

    def prep_tables(self, document, **dargs):
        """Prepare tables."""
        for key, item in dargs.get('value').items():
            index = dargs.get('index')
            row = (index % 9) / 3
            col = index % 3

            if (index % 9 == 0 and index > 0):
                table = document.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

            document = self.build_tables(
                document, table=dargs.get('table'), index=index, row=row,
                col=col, item=item, key=key, graphics=dargs.get('graphics'))
            return document

    def build_tables(self, document, **dargs):
        """Construct tables."""
        settings = self.settings
        index = dargs.get('index')
        row = dargs.get('row')
        col = dargs.get('col')
        item = dargs.get('item')
        table = dargs.get('table')

        if (index % 18 == 0 and index > 17):
            document.add_page_break()
            table = document.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

        paragraph = table.cell(int(row), int(col)).paragraphs[0]
        paragraph.paragraph_format.line_spacing = 0.0
        if dargs.get('graphics'):
            table.cell(int(row), int(col)).add_picture(
                'static/profiles/{}/img/clients/{}.png'.format(
                    settings.DIR,
                    dargs.get('key')
                ),
                width=self.cm(4.8)
            )
        else:
            table.cell(int(row), int(col)).add_paragraph('')
        document = self.finish_tables(
            document, table=table, row=row, col=col, item=item,
            key=dargs.get('key'))
        return document

    def finish_tables(self, document, **dargs):
        """Complete process started in calling method."""
        table = dargs.get('table')
        row = int(dargs.get('row'))
        col = int(dargs.get('col'))
        self.projects = dargs.get('item').get('projects')

        current_cell = table.cell(row, col)
        current_cell.width = self.cm(5)
        paragraph = table.cell(row, col).paragraphs[1]
        paragraph.paragraph_format.space_after = 0

        paragraph = table.cell(row, col).add_paragraph(
            dargs.get('item').get('position'), style='Heading 4')
        paragraph.paragraph_format.line_spacing = self.spacing
        paragraph.paragraph_format.space_after = 0

        paragraph = table.cell(row, col).add_paragraph(
            dargs.get('item').get('company'), style='Heading 5')
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = 0

        paragraph = table.cell(row, col).add_paragraph(
            "{0}, {1}".format(
                dargs.get('item').get('location'),
                dargs.get('item').get('duration')),
            style='Heading 6')
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = 0
        document = self.add_projects(
            document, table, row, col)
        return document

    def add_projects(self, document, table, row, col):
        """Add projects to experience section.

        :param [dict(str, str)] projects:
            Projects for a portion of Experience section.
        :param object table: Table from current document.
        :param int row: Index of current row.
        :param int col: Index of current col.
        :return: Updated Projects table.
        :rtype: object
        """
        for project_name, project in self.projects.items():
            paragraph = table.cell(row, col).add_paragraph(
                project_name.replace('_', ' ').title(),
                style='List Bullet')
            paragraph.paragraph_format.line_spacing = self.spacing
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.tab_stops.add_tab_stop(self.cm(0.2))
            for item in project:
                item_paragraph = table.cell(row, col).add_paragraph(
                    item,
                    style='List Bullet 2')
                item_paragraph.paragraph_format.line_spacing = 0.8
                item_paragraph.paragraph_format.space_after = 0
                item_paragraph.paragraph_format.tab_stops.add_tab_stop(
                    self.cm(10))
        return document
