"""Experience module."""
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_projects(projects, table, row, col):
    """Add projects to experience section.

    :param [dict(str, str)] projects:
        Projects for a portion of Experience section.
    :param object table: Table from current document.
    :param int row: Index of current row.
    :param int col: Index of current col.
    :return: Updated Projects table.
    :rtype: object
    """
    for name, project in projects.items():
        paragraph = table.cell(row, col).add_paragraph(
            name.replace('_', ' ').title(),
            style='List Bullet')
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = 0
        for item in project:
            paragraph = table.cell(row, col).add_paragraph(
                item,
                style='List Bullet 2')
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = 0
    return table


def add_experience(settings, sections, document):
    """Add experience section.

    :param [dict(str, str)] experience:
        Experience section to add to document.
    :param objectable document: Documentable to update.
    :return: Documentable updated with Experience section.
    :rtype: object
    """
    experience = sections[4].get('experience')
    paragraph = document.add_paragraph('')
    paragraph.paragraph_format.line_spacing = 0.0
    paragraph.paragraph_format.page_break_before = True
    paragraph = document.add_paragraph(
        'Experience', style='Heading 3')
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(experience):
        if index % 3 == 0:
            table.add_row()
        document = set_tables(
            document=document, table=table, value=value,
            index=index, settings=settings)
    paragraph = document.add_paragraph('')
    paragraph.paragraph_format.line_spacing = 0
    return document


def set_tables(document, **dargs):
    """Set tables for the experience section."""
    document = (
        prep_tables(
            document=document, table=dargs.get('table'),
            value=dargs.get('value'), index=dargs.get('index'),
            settings=dargs.get('settings'))
    )
    return document


def prep_tables(document, **dargs):
    """Prepare tables."""
    for item in dargs.get('value'):
        index = dargs.get('index')
        settings = dargs.get('settings')
        row = (index % 9) / 3
        col = index % 3
        if (
                index % 9 == 0 and
                index > 0
        ):
            table = document.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

        document = build_tables(
            document=document, table=dargs.get('table'),
            settings=settings, index=index,
            row=row, col=col, item=item)
        return document


def build_tables(document, **dargs):
    """Construct tables."""
    index = dargs.get('index')
    row = dargs.get('row')
    col = dargs.get('col')
    settings = dargs.get('settings')
    item = dargs.get('item')
    table = dargs.get('table')
    if (index % 18 == 0 and index > 17):
        document.add_page_break()
        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    paragraph = table.cell(row, col).paragraphs[0]
    paragraph.paragraph_format.line_spacing = 0.0
    table.cell(row, col).add_picture(
        '/srv/rsum/static/{}/img/970x647/{}.jpg'.format(
            settings.DIR,
            index+1
        ),
        width=Cm(4.8)
    )
    document = finish_tables(
        document=document, table=table, row=row, col=col, item=item)
    return document


def finish_tables(document, **dargs):
    """Complete process started in calling methond."""
    table = dargs.get('table')
    row = dargs.get('row')
    col = dargs.get('col')
    item = dargs.get('item')
    paragraph = table.cell(row, col).paragraphs[1]
    paragraph.paragraph_format.space_after = 0

    paragraph = table.cell(row, col).add_paragraph(
        item, style='Heading 4')
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = 0
    paragraph = table.cell(row, col).add_paragraph(
        item, style='Heading 5')
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = 0
    paragraph = table.cell(row, col).add_paragraph(
        "{0}, {1}".format(item, item),
        style='Heading 6')
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = 0
    table = add_projects(item.get('projects'), table, row, col)
    return document
