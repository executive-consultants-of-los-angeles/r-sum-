"""Skills section module."""
# pylint: disable=no-name-in-module
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_paragraph(t_sub, index):
    """Add skills section.

    :param t_sub: Table for the current sub skill.
    :param index: Index of current skill.
    """
    paragraph = t_sub.cell(index-1, 0).paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = 0

    return paragraph


def set_inner_paragraph(t_sub, index):
    """Set styles for inner skills section.

    :param t_sub: Table for the current sub skill.
    :param index: Index of current skill.
    """
    paragraph = t_sub.cell(index-1, 1).paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return paragraph


class Skills(object):
    """Class for skills section object."""

    name = None
    graphics = False
    document = None
    sub_skills = None
    current_year = float(datetime.datetime.now().strftime('%Y'))

    def save(self, section, document, graphics):
        """Add skills section.

        :param skills: Skills section to add to document.
        :type summary: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Skills.
        :rtype: object
        """
        self.document = document
        self.graphics = graphics
        document = self.get_skills(section, document)
        return document

    def get_skills(self, section, document):
        """Add skills section.

        :param skills: Skills section to add to document.
        :type summary: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Skills.
        :rtype: object
        """
        skills = section

        document.add_paragraph('Skills', style='Heading 3')
        table = document.add_table(rows=1, cols=2)
        index = 1
        for skill_name, skill in skills.items():
            if isinstance(skill, dict):
                experience = int(self.current_year) - int(skill.get('start'))
                experience = '{0} year(s)'.format(str(experience))
                output_name = skill_name.replace('_', ' ').title()
                table.cell(0, index % 2).add_paragraph(
                    "{} \t || {}".format(output_name, experience),
                    "List Bullet"
                )

                for sub in skill.items():
                    if isinstance(sub[1], dict):
                        experience = (
                            int(self.current_year) - int(sub[1].get('start')))
                        experience = '{0} year(s)'.format(str(experience))
                        table.cell(0, index % 2).add_paragraph(
                            "{} \t || {}".format(
                                sub[1].get('name'), experience),
                            "List Bullet 2"
                        )
            index = index + 1
        return document
