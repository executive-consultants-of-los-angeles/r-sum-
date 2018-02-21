"""Summary section export."""
# pylint: disable=no-member
from django.conf import settings as django_settings
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


class Summary(object):
    """Sumary section object class.

    .. attribute:: name

       Name of section, deprecated.

    .. attribute:: settings

       Django settings.

    .. attribute:: cm

       Cm measurement.

    .. attribute:: summary_table

       Ostensibly the table in which to display the summary.
       Most likely depcrecated.
    """

    name = None
    settings = django_settings
    cm = Cm
    summary_table = None

    def save(self, section, document, graphics):
        """Add summary section.

        :param dict summary: Summary section to add to document.
        :param docx.Document document: Currenlty active document.
        :param bool graphics: True to load graphics, false to not.
        :return: Document updated with Summary with or without graphics.
        """
        if graphics:
            document = self.get_summary_graphics(section, document)
        else:
            document = self.get_summary(section, document)
        return document

    def get_summary_graphics(self, section, document):
        """Acquire the graphics for insertion into the current document.

        :param dict section: Summary section to add to document.
        :param docx.Document document: Currently active document.
        :return: Document updated with Summary and pretty pictures.
        """
        summary = section
        settings = self.settings
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0

        summary_table = document.add_table(rows=1, cols=2)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        summary_table.cell(0, 0).width = Cm(6)
        summary_table.cell(0, 0).add_picture(
            'static/profiles/{0}/img/500x700/02.png'.format(
                settings.DIR),
            width=Cm(5))

        summary_table.cell(0, 1).add_paragraph(
            'Summary', style='Heading 3')
        summary_table.cell(0, 1).add_paragraph(
            summary.get('content'), style='Normal')
        self.summary_table = summary_table

        paragraph = self.format_paragraph(
            summary_table.cell(0, 1).paragraphs[1])

        return document

    def get_summary(self, section, document):
        """Add summary section.

        :param dict section: Summary section to add to document.
        :param docx.Document document: Current document.
        :return: Document updated with Summary.
        """
        summary = section
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0

        summary_table = document.add_table(rows=1, cols=1)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        summary_table.cell(0, 0).add_paragraph(
            'Summary', style='Heading 3')
        summary_table.cell(0, 0).add_paragraph(
            summary.get('content'), style='Normal')
        self.summary_table = summary_table

        paragraph = self.format_paragraph(
            summary_table.cell(0, 0).paragraphs[1])

        return document

    def format_paragraph(self, paragraph):
        """Accept a paragraph, format it, then return it.

        :param docx.Paragraph paragraph: Paragraph to format.
        :return: Formatted paragraph.
        """
        paragraph.paragraph_format.line_spacing = 1.0
        self.summary_table.cell(0, 0).width = self.cm(10)
        paragraph = self.summary_table.cell(0, 0).paragraphs[0]
        paragraph.paragraph_format.line_spacing = 0.0
        return paragraph
