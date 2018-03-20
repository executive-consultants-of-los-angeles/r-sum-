"""Contact module."""
# pylint: disable=too-few-public-methods
from docx.shared import Pt


def complete_contact(contact, document, table):
    """Complete addition of contact section."""
    table.cell(0, 1).add_paragraph(
        'Email',
        style='Heading 4')
    paragraph = table.cell(0, 1).add_paragraph(
        contact.get('email'),
        style='Heading 5')
    paragraph.paragraph_format.space_before = 0
    paragraph = table.cell(1, 0).paragraphs[0]
    paragraph.paragraph_format.line_spacing = Pt(7)
    table.cell(1, 0).add_paragraph(
        'Phone',
        style='Heading 4')
    paragraph = table.cell(1, 0).add_paragraph(
        contact.get('phone'),
        style='Heading 5')
    paragraph.paragraph_format.space_before = 0
    paragraph = table.cell(1, 1).paragraphs[0]
    paragraph.paragraph_format.line_spacing = Pt(7)
    table.cell(1, 1).add_paragraph('Location', style='Heading 4')
    paragraph = table.cell(1, 1).add_paragraph(
        contact.get('location'),
        style='Heading 5')
    paragraph.paragraph_format.space_before = 0
    return document


def start_contact(contact, document):
    """Start recording contact."""
    paragraph = document.add_paragraph('')
    paragraph.paragraph_format.line_spacing = 0.0
    document.add_paragraph(
        contact.get('title'), style='Heading 3')

    paragraph = document.add_paragraph(
        contact.get('message'), style='Heading 6')

    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = Pt(7)
    table = document.add_table(rows=2, cols=2)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.line_spacing = Pt(7)
    table.cell(0, 0).add_paragraph(
        'Email',
        style='Heading 4')
    paragraph = table.cell(0, 0).add_paragraph(
        contact.get('email'),
        style='Heading 5')
    paragraph.paragraph_format.space_before = 0
    paragraph = table.cell(0, 1).paragraphs[0]
    paragraph.paragraph_format.line_spacing = Pt(7)
    # document = complete_contact(contact, document, table)
    return document


class Contact(object):
    """Class for contact object export."""

    name = None

    def save(self, section, document, graphics):
        """Add contact section.

        :param [dict(str, str)] contact:
            Contact section for current document.
        :param object document: Current document.
        :return: Current document with Contact section.
        :rtype: object
        """
        contact = section
        self.name = graphics
        document = start_contact(contact, document)
        return document
