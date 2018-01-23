"""Skills section module."""
# pylint: disable=no-name-in-module
import datetime
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Skills(object):
    """Class for skills section object."""

    name = None
    document = None
    sub_skills = None

    def save(self, name, section, document):
        """Add skills section.

        :param skills: Skills section to add to document.
        :type summary: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Skills.
        :rtype: object
        """
        self.name = name
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0
        current_year = datetime.datetime.now().strftime("%Y")
        skills = section

        table = document.tables[1]
        table.cell(0, 1).add_paragraph('Skills', style='Heading 3')
        t_sub = table.cell(0, 1).add_table(rows=1, cols=2)
        table.cell(0, 1).tables[0].columns[0].width = Cm(7)
        index = 1
        for skill_name, skill in skills.items():
            if isinstance(skill, dict):
                experience = int(current_year) - int(skill.get('start'))
                experience = '{0} year(s)'.format(str(experience))
                output_name = skill_name.replace('_', ' ').title()

                # Add a row to the sub table.
                t_sub.add_row()
                t_sub.cell(index-1, 0).textable = output_name
                paragraph = t_sub.cell(index-1, 0).paragraphs[0]
                paragraph.style = 'Skill'
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = 0
                t_sub.cell(index-1, 1).textable = experience

                paragraph = t_sub.cell(index-1, 1).paragraphs[0]
                paragraph.style = 'Skill'
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = 0
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                t_sub = self.add_sub_skills(skill, t_sub, index-1)
            index = index + 1
        return document

    def add_sub_skills(self, subs, skilltable, skilltable_index):
        """Add sub skills to skills section.

        :param subs: Sub skills to add to document.
        :type subs: [dict(str, str)]
        :param object ts: Table cell to update.
        :param int ts_index: Index for current table cell.
        :return: Document updated with sub skills.
        :rtype: object
        """
        self.sub_skills = subs
        current_year = float(datetime.datetime.now().strftime("%Y"))
        sub_table = skilltable.cell(
            skilltable_index, 0).add_table(rows=1, cols=2)
        index = 0
        for sub in subs:
            if isinstance(sub, dict):
                experience = int(current_year) - int(sub.get('start'))
                experience = '{0} year(s)'.format(str(experience))
                if index == 0:
                    sub_table.cell(0, 0).text = sub.get('name')
                    sub_table.cell(0, 1).text = experience
                    sub_table.cell(0, 0).width = Cm(5)
                else:
                    sub_table.add_row()
                    sub_table.cell(index, 0).text = sub.get('name')
                    sub_table.cell(index, 0).width = Cm(1)
                    sub_table.cell(index, 1).text = experience
                paragraph = sub_table.cell(index, 0).paragraphs[0]
                paragraph.style = 'Sub Skill'
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = 0
                paragraph = sub_table.cell(index, 1).paragraphs[0]
                paragraph.style = 'Sub Skill'
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph.paragraph_format.space_after = 0
                index = index + 1
        return skilltable
