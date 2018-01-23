"""Module for education."""
# pylint: disable=no-name-in-module,no-member
from django.conf import settings as django_settings
from docx.shared import Cm


class Education(object):
    """Education export object."""

    name = None
    document = None
    settings = django_settings

    def save(self, name, section, document):
        """Add education section.

        :param [dict(str, str)] education:
            Education section for current document.
        :param object document: Current document.
        :return: Current document with Educaiton section.
        :rtype: object
        """
        education = section
        self.name = name
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0
        settings = self.settings
        paragraph = document.add_paragraph(
            'Education',
            style='Heading 3')
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.page_break_before = True
        document.add_picture(
            '../static/profiles/{0}/img/1920x1080/01.jpg'.format(
                settings.DIR),
            width=Cm(4))
        paragraph = document.add_paragraph(
            education.get('name'),
            style='Heading 4')
        paragraph.paragraph_format.space_before = 0
        paragraph = document.add_paragraph(
            education.get('studies'),
            style='Heading 5')
        paragraph.paragraph_format.space_before = 0
        paragraph = document.add_paragraph(
            "{0}, {1}".format(
                education.get('location'),
                education.get('duration')),
            style='Heading 6')
        paragraph.paragraph_format.space_before = 0
        print(name)
        print(document)
        return document

    def organize(self):
        """Organize something."""
        return self
