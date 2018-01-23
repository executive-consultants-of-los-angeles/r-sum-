"""Summary section export."""
# pylint: disable=no-member
from django.conf import settings as django_settings
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.text import WD_ALIGN_PARAGRAPH


class Summary(object):
    """Sumary section object class."""

    name = None
    settings = django_settings

    def save(self, name, section, document):
        """Add summary section.

        :param summary: Summary section to add to document.
        :type summary: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Summary.
        :rtype: object
        """
        settings = self.settings
        self.name = name
        summary = section
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.line_spacing = 0.0

        summary_table = document.add_table(rows=1, cols=2)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        summary_table.cell(0, 0).width = Cm(6)
        summary_table.cell(0, 0).add_picture(
            '../static/profiles/{0}/img/500x700/02.png'.format(settings.DIR),
            width=Cm(5))

        summary_table.cell(0, 1).add_paragraph('Summary', style='Heading 3')
        summary_table.cell(0, 1).add_paragraph(
            summary.get('content'), style='Normal')
        paragraph = summary_table.cell(0, 1).paragraphs[1]
        paragraph.paragraph_format.line_spacing = 1.0
        summary_table.cell(0, 1).width = Cm(10)
        paragraph = summary_table.cell(0, 0).paragraphs[0]
        paragraph.paragraph_format.line_spacing = 0.0
        paragraph = summary_table.cell(0, 1).paragraphs[0]
        paragraph.paragraph_format.line_spacing = 0.0
        print(name)
        print(document)
        return document

    def no_action(self):
        """Do nothing."""
        return self
